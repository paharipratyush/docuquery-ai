"""
DocuQuery AI
----------------------------
Highlighting document-based queries with AI-powered responses.
Supports multiple document formats including PDF, DOCX, TXT, and URLs.
"""

import streamlit as st
import faiss
import os
import time
import validators
import traceback
from typing import List, Union, Optional
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_huggingface import HuggingFaceEndpoint
from streamlit.runtime.uploaded_file_manager import UploadedFile
from tqdm import tqdm

# Initialize environment and configurations
try:
    from secret_api_keys import huggingface_api_key
    os.environ['HUGGINGFACEHUB_API_TOKEN'] = huggingface_api_key
except ImportError:
    st.error("Error: secret_api_keys.py not found. Please create this file with your huggingface_api_key.")
    st.stop()

# Constants
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
CHUNK_SIZE = 2000
CHUNK_OVERLAP = 200
URL_TIMEOUT = 30  # seconds

class DocumentProcessor:
    """Handles document processing for various input types."""

    def __init__(self):
        self.embeddings = self._initialize_embeddings()

    @staticmethod
    def _initialize_embeddings():
        """Initialize and return the embedding model."""
        model_name = "sentence-transformers/all-mpnet-base-v2"
        return HuggingFaceEmbeddings(
            model_name=model_name,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': False}
        )

    def validate_url(self, url: str) -> bool:
        """Validate if the given URL is properly formatted."""
        return validators.url(url)

    def process_urls(self, urls: List[str]) -> List[str]:
        """Process URLs and extract text content."""
        texts = []
        for url in tqdm(urls, desc="Processing URLs"):
            if url.strip():  # Only process non-empty URLs
                try:
                    loader = WebBaseLoader(url)
                    loader.requests_kwargs = {"timeout": URL_TIMEOUT}
                    documents = loader.load()
                    texts.extend([doc.page_content for doc in documents])
                except Exception as e:
                    st.warning(f"Error processing URL {url}: {str(e)}")
        return texts

    def process_pdf(self, file: Union[BytesIO, UploadedFile]) -> str:
        """Process PDF file and extract text."""
        if isinstance(file, UploadedFile) and file.size > MAX_FILE_SIZE:
            raise ValueError(f"File size exceeds limit of {MAX_FILE_SIZE/1024/1024}MB")

        try:
            if isinstance(file, UploadedFile):
                pdf_reader = PdfReader(BytesIO(file.read()))
            else:
                pdf_reader = PdfReader(file)

            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Error processing PDF: {str(e)}")

    def process_docx(self, file: Union[BytesIO, UploadedFile]) -> str:
        """Process DOCX file and extract text."""
        try:
            if isinstance(file, UploadedFile):
                doc = Document(BytesIO(file.read()))
            else:
                doc = Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            raise ValueError(f"Error processing DOCX: {str(e)}")

    def process_txt(self, file: Union[BytesIO, UploadedFile]) -> str:
        """Process TXT file and extract text."""
        try:
            if isinstance(file, UploadedFile):
                return file.read().decode("utf-8")
            else:
                return file.read().decode("utf-8")
        except Exception as e:
            raise ValueError(f"Error processing TXT: {str(e)}")

    def create_vectorstore(self, texts: List[str], existing_vectorstore: Optional[FAISS] = None) -> FAISS:
        """Create or update a FAISS vectorstore from the given texts."""
        try:
            text_splitter = CharacterTextSplitter(
                separator="\n",
                chunk_size=CHUNK_SIZE,
                chunk_overlap=CHUNK_OVERLAP,
                length_function=len
            )

            split_texts = text_splitter.split_text(" ".join(texts))

            if not split_texts:
                raise ValueError("No text chunks generated after splitting")

            if existing_vectorstore:
                existing_vectorstore.add_texts(split_texts)
                return existing_vectorstore
            else:
                dimension = 768  # HuggingFace embedding dimension
                index = faiss.IndexFlatL2(dimension)

                vector_store = FAISS(
                    embedding_function=self.embeddings.embed_query,
                    index=index,
                    docstore=InMemoryDocstore(),
                    index_to_docstore_id={}
                )

                vector_store.add_texts(split_texts)
                return vector_store

        except Exception as e:
            raise ValueError(f"Error creating vector store: {str(e)}")

def reset_session_state():
    """Reset all session state variables."""
    for key in list(st.session_state.keys()):
        del st.session_state[key]

def answer_question(vectorstore: FAISS, query: str) -> dict:
    """Generate answer for the given query using the vectorstore."""
    try:
        llm = HuggingFaceEndpoint(
            endpoint_url="https://api-inference.huggingface.co/models/meta-llama/Meta-Llama-3-8B-Instruct",
            huggingface_api_key=huggingface_api_key,
            max_new_tokens=256,
            temperature=0.7,
            model_kwargs={"max_length": 7936}
        )

        qa = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vectorstore.as_retriever(
                search_kwargs={"k": 3}
            )
        )

        return qa({"query": query})
    except Exception as e:
        raise ValueError(f"Error generating answer: {str(e)}")

def main():
    st.set_page_config(page_title="DocuQuery AI", layout="wide")
    st.title("📚 DocuQuery AI")

    # Initialize processor
    doc_processor = DocumentProcessor()

    # Sidebar
    with st.sidebar:
        st.header("Settings")
        if st.button("Reset Session"):
            reset_session_state()
            st.success("Session reset successfully!")

    # Main input selection
    input_type = st.selectbox("Select Input Type", ["Link", "PDF", "Text", "DOCX", "TXT"])

    try:
        input_data = None
        if input_type == "Link":
            number_input = st.number_input(
                "Number of URLs",
                min_value=1,
                max_value=20,
                value=1,
                step=1
            )
            input_data = []
            for i in range(int(number_input)):
                url = st.text_input(f"URL {i+1}", key=f"url_{i}")
                if url:
                    input_data.append(url)

        elif input_type == "Text":
            input_data = st.text_area("Enter Text", height=200)

        elif input_type in ["PDF", "DOCX", "TXT"]:
            file_type = {"PDF": ["pdf"], "DOCX": ["docx", "doc"], "TXT": ["txt"]}[input_type]
            uploaded_file = st.file_uploader(f"Upload {input_type} file", type=file_type)
            if uploaded_file:
                if uploaded_file.size > MAX_FILE_SIZE:
                    st.error(f"File size exceeds maximum limit of {MAX_FILE_SIZE/1024/1024}MB")
                else:
                    input_data = uploaded_file

        # Process button
        if st.button("Proceeed", disabled=not input_data):
            try:
                with st.spinner("Processing input..."):
                    if input_type == "Link":
                        texts = doc_processor.process_urls(input_data)
                    elif input_type == "PDF":
                        texts = [doc_processor.process_pdf(input_data)]
                    elif input_type == "DOCX":
                        texts = [doc_processor.process_docx(input_data)]
                    elif input_type == "TXT":
                        texts = [doc_processor.process_txt(input_data)]
                    else:
                        texts = [input_data]

                    vectorstore = doc_processor.create_vectorstore(
                        texts,
                        existing_vectorstore=st.session_state.get("vectorstore")
                    )
                    st.session_state["vectorstore"] = vectorstore
                    st.success("Processing complete! You can now ask questions.")

            except Exception as e:
                st.error(f"Error during processing: {str(e)}")
                st.error(traceback.format_exc())

        # Question answering section
        if "vectorstore" in st.session_state:
            st.header("Ask Questions")
            query = st.text_input("Enter your question:")

            if st.button("Get Answer", disabled=not query):
                try:
                    with st.spinner("Generating answer..."):
                        answer = answer_question(st.session_state["vectorstore"], query)
                        st.write("Answer:", answer["result"])
                except Exception as e:
                    st.error(f"Error generating answer: {str(e)}")

    except Exception as e:
        st.error(f"An unexpected error occurred: {str(e)}")
        st.error(traceback.format_exc())

if __name__ == "__main__":
    main()
